import json
import re
from docx import Document
from typing import List, Dict, Optional
import openai


def extract_text_from_docx(docx_file) -> str:
    try:
        document = Document(docx_file) if hasattr(docx_file, "read") else Document(str(docx_file))
        return "\n".join([para.text for para in document.paragraphs])
    except Exception as e:
        return f"Error: {e}"

def extract_invoice_data_from_docx(docx_file) -> dict:
    """
    Extract invoice-related data from a Word (.docx) document.
    :param docx_file: Path to the .docx file or a file-like object
    :return: Dictionary with extracted invoice fields
    """
    try:
        if hasattr(docx_file, "read"):
            document = Document(docx_file)
        else:
            document = Document(str(docx_file))
    except Exception as e:
        return {"error": f"Failed to open Word document: {e}"}

    text = "\n".join([para.text for para in document.paragraphs])

    # Basic regex-based parsing (customize as needed)
    invoice_data = {
        "invoice_id": _extract_field(text, r"Invoice\s*(Number|No\.?):?\s*([A-Za-z0-9\-]+)"),
        "date": _extract_field(text, r"Date[:\s]*([\d/.-]+)"),
        "vendor": _extract_field(text, r"Vendor[:\s]*([A-Za-z\s&]+)"),
        "amount": _extract_field(text, r"(Total|Amount Due|Amount)[:\s]*\$?([\d,]+\.\d{2})"),
        "raw_text": text  # for debugging or fallback
    }

    return invoice_data

# utils/business_rules.py

def extract_business_rules_from_docx(file_path: str) -> List[Dict[str, str]]:
    doc = Document(file_path)
    rules = []
    rule_number = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name in ["List Number", "Heading 2", "Normal"]:
            if text.lower().startswith("rule") or text[:2].isdigit():
                rules.append({
                    "rule_number": f"Rule {rule_number}",
                    "description": text
                })
                rule_number += 1

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                rule_text = cells[1].text.strip()
                if rule_text:
                    rules.append({
                        "rule_number": cells[0].text.strip() or f"Rule {rule_number}",
                        "description": rule_text
                    })
                    rule_number += 1

    return rules


def extract_business_rules_from_docx_with_confidence_score(file_path: str, llm_model=None) -> List[Dict[str, str]]:
    """
    Extracts business rules from a .docx file, including both paragraphs and tables.
    Optionally uses LLM to assign confidence scores.

    Args:
        file_path (str): Path to the .docx file.
        llm_model: Optional LLM model to assign confidence scores.

    Returns:
        List[Dict[str, str]]: Extracted rules with optional confidence.
    """
    doc = Document(file_path)
    rules = []
    rule_number = 1

    # Extract from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name in ["List Number", "Heading 2", "Normal"]:
            if text.lower().startswith("rule") or text[:2].isdigit():
                rule = {
                    "rule_number": f"Rule {rule_number}",
                    "description": text
                }
                if llm_model:
                    rule["confidence"] = get_confidence_score_llm(llm_model, text)
                rules.append(rule)
                rule_number += 1

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                rule_text = cells[1].text.strip()
                if rule_text:
                    rule = {
                        "rule_number": cells[0].text.strip() or f"Rule {rule_number}",
                        "description": rule_text
                    }
                    if llm_model:
                        rule["confidence"] = get_confidence_score_llm(llm_model, rule_text)
                    rules.append(rule)
                    rule_number += 1

    return rules

def _extract_field(text: str, pattern: str) -> str:
    match = re.search(pattern, text, re.IGNORECASE)
    if match:
        return match.group(len(match.groups()))
    return ""

def get_confidence_score_llm(llm_model, rule_text: str) -> str:
    """
    Placeholder function for confidence score using an LLM.

    Args:
        llm_model: An LLM model or chain
        rule_text: The text of the rule

    Returns:
        A string confidence score, e.g. '92%'
    """
    # Example only — replace with real model call
    response = llm_model.predict(f"What is the confidence that this is a valid business rule? \"{rule_text}\"")
    return response.strip()  # Should return '95%' or similar

def extract_business_rules_from_docx_basic(file_path: str) -> list[str]:
    doc = Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]


