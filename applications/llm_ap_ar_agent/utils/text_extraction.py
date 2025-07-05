import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from PyPDF2 import PdfReader
import logging
import pdfplumber
from utils.pdf_utils import extract_text_from_scanned_pdf

logger = logging.getLogger(__name__)
def extract_amount(text):
    # Priority: $amount pattern
    match = re.search(r"\$\s*(\d+(?:,\d{3})*(?:\.\d{2})?)", text)
    if match:
        return int(match.group(1).replace(",", ""))

    # Fallback: number after 'for' or 'amount'
    match = re.search(r"(?:for|amount)\s*(\d+(?:,\d{3})*)", text, re.IGNORECASE)
    if match:
        return int(match.group(1).replace(",", ""))

    # General fallback: any large number
    match = re.search(r"\b(\d{3,})\b", text)
    if match:
        return int(match.group(1).replace(",", ""))

    return None

def regex_extract_fields(text):
    result = {}
    patterns = {
        "invoice_id": r"Invoice ID[:\s]*([A-Z0-9\-]+)",
        "vendor": r"Vendor[:\s]*(.+)",
        "amount": r"Amount[:\s]*\$?([\d,.]+)",
        "date": r"Date[:\s]*([0-9]{4}-[0-9]{2}-[0-9]{2})",
        "po_number": r"PO[:\s]*([A-Z0-9\-]+)"
    }
    for field, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if field == "amount":
                value = float(value.replace(",", ""))
            result[field] = value
    return result

def safe_json_parse(output: str):
    try:
        start = output.find("{")
        end = output.rfind("}") + 1
        json_str = output[start:end]
        return json.loads(json_str)
    except Exception:
        # Fallback: regex key-value parsing
        data = {}
        kv_pairs = re.findall(r"(\w+)\s*=\s*([^\n]+)", output)
        for key, value in kv_pairs:
            data[key.strip()] = value.strip()
        return {
            "warning": "Parsed from plain text, not valid JSON",
            "data": data,
            "raw_output": output
        }

def extract_text_from_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        return f"[Error extracting PDF text: {str(e)}]"

def extract_invoice_table(pdf_path: str) -> list:
    """Attempt to extract tabular data (e.g., invoice line items) from PDF."""
    try:
        tables = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_table()
                if extracted:
                    tables.append(extracted)
        return tables
    except Exception as e:
        logger.warning(f"[PDF Table Extract] Failed: {e}")
        return []

def robust_extract_text(pdf_path: str) -> str:
    """Try digital extraction first, fall back to OCR if needed."""
    text = extract_text_from_pdf(pdf_path)
    if not text or len(text.strip()) < 100:
        logger.info("[PDF Extract] Falling back to OCR...")
        text = extract_text_from_scanned_pdf(pdf_path)
    return text.strip()

def extract_text_from_image(image_file):
    try:
        image = Image.open(image_file)
        return pytesseract.image_to_string(image).strip()
    except Exception as e:
        return f"[Error extracting image text: {str(e)}]"

def extract_text_from_image_embedded_pdf(pdf_file):
    try:
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        text = ""

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Attempt to extract text directly
            text += page.get_text()

            # OCR fallback for embedded images
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                image = Image.open(io.BytesIO(image_bytes))
                ocr_text = pytesseract.image_to_string(image)
                text += f"\n[Image OCR Page {page_num + 1} - Image {img_index + 1}]\n{ocr_text}\n"

        return text.strip()
    except Exception as e:
        return f"[Error extracting image-embedded PDF text: {str(e)}]"

def extract_rules_from_docx(file) -> str:
    """
    Extracts rules from a .docx file and returns a structured JSON-like string.
    If no structured pattern is found, returns raw text.
    """

    if isinstance(file, BytesIO):
        doc = Document(file)
    else:
        doc = Document(BytesIO(file.read()))

    rules = {}
    raw_lines = []

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        raw_lines.append(line)

        # Simple pattern: "Rule Name: Rule Description"
        if ':' in line:
            key, value = line.split(':', 1)
            rules[key.strip()] = value.strip()

    if rules:
        return json.dumps(rules, indent=2)
    else:
        # Fallback to plain text
        return "\n".join(raw_lines)

from docx import Document
from io import BytesIO
import json

def infer_type(value: str):
    """Try to cast string value to int, float, or bool."""
    value = value.strip()
    lowered = value.lower()
    if lowered in ['true', 'yes']:
        return True
    elif lowered in ['false', 'no']:
        return False
    try:
        if '.' in value:
            return float(value)
        return int(value)
    except ValueError:
        return value

def extract_rules_from_docx_with_type_inference(file) -> dict:
    """
    Extracts business rules from a DOCX file and infers types.
    Returns a dictionary of rules.
    """

    if isinstance(file, BytesIO):
        doc = Document(file)
    else:
        doc = Document(BytesIO(file.read()))

    rules = {}
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line or ':' not in line:
            continue
        key, value = line.split(':', 1)
        rules[key.strip()] = infer_type(value)

    return rules

def extract_business_rules_from_docx(file_path: str) -> list[str]:
    doc = Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

import pandas as pd
from typing import List, Dict, Any

import pandas as pd
from typing import Dict, Any
import re

def extract_timecard_from_excel(uploaded_file: str, sheet_name: str) -> str:
    timecard = pd.read_excel(uploaded_file, sheet_name=None, engine='openpyxl', usecols='A,B,K,L')
    timecard_sheets = timecard.keys()
    logger.debug(timecard_sheets)
    return timecard.get(sheet_name).to_string()

def extract_timecard_metadata_generic(file_path: str) -> Dict[str, Any]:
    df = pd.read_excel(file_path, header=None)
    metadata = {}

    # Known field patterns (regex-style)
    field_patterns = {
        "employee_name": r"\bemployee\b",
        "manager_name": r"\bmanager\b",
        "street_address": r"street address",
        "address_2": r"address 2",
        "city_state_zip": r"city.*zip",
        "phone": r"phone",
        "email": r"e-?mail",
        "week_ending": r"week ending",
        "total_hours": r"total hours?",
        "hourly_rate": r"rate per hour",
        "total_amount": r"total pay|total amount",
    }

    def match_field(label: str) -> str:
        label_lower = label.strip().lower()
        for field, pattern in field_patterns.items():
            if re.search(pattern, label_lower):
                return field
        return None

    # Scan row-by-row
    for _, row in df.iterrows():
        row = row.fillna("").astype(str).str.strip().tolist()
        for i in range(len(row) - 1):
            field_name = match_field(row[i])
            if field_name and row[i + 1]:
                value = row[i + 1]
                # Try converting numerical fields
                if field_name in ["total_hours", "hourly_rate", "total_amount"]:
                    try:
                        value = float(re.sub(r"[^\d.]", "", value))
                    except:
                        value = 0.0
                metadata[field_name] = value

    # Build full address
    full_address = ", ".join(filter(None, [
        metadata.get("street_address", ""),
        metadata.get("address_2", ""),
        metadata.get("city_state_zip", "")
    ]))

    return {
        "employee_name": metadata.get("employee_name", "Unknown"),
        "manager_name": metadata.get("manager_name", "Unknown"),
        "phone": metadata.get("phone", "Unknown"),
        "email": metadata.get("email", "Unknown"),
        "address": full_address or "Unknown",
        "week_ending": metadata.get("week_ending", "Unknown"),
        "total_hours": metadata.get("total_hours", 0.0),
        "hourly_rate": metadata.get("hourly_rate", 0.0),
        "total_amount": metadata.get("total_amount", 0.0),
    }

def create_icl_prompt(extracted_text: str):
    return f"""
Example 1:
Invoice Text:
Invoice ID: INV-1234
Vendor: Acme Corp
Amount: $1234.56
Date: 2025-04-15
PO: PO-9876

Extracted JSON:
{{
  "invoice_id": "INV-1234",
  "vendor": "Acme Corp",
  "amount": 1234.56,
  "date": "2025-04-15",
  "po_number": "PO-9876"
}}

Example 2:
Invoice Text:
This invoice is issued by Globex Inc on 2025-01-10.
Invoice ID: INV-5678
The amount due is $7890.00. PO: PO-1122

Extracted JSON:
{{
  "invoice_id": "INV-5678",
  "vendor": "Globex Inc",
  "amount": 7890.00,
  "date": "2025-01-10",
  "po_number": "PO-1122"
}}

Now extract JSON for this:
Invoice Text:
{extracted_text}

Respond with JSON only.
"""


